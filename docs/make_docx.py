"""Build the BuildSprint idea briefs as a formatted .docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)
CODE_BG = "F2F2F2"

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for name, size, color in [("Heading 1", 18, ACCENT), ("Heading 2", 14, ACCENT), ("Heading 3", 12, GREY)]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def para(text="", bold_prefix=None, size=11, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(size)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)


def codeblock(lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade(cell, CODE_BG)
    cell.width = Inches(6.0)
    p = cell.paragraphs[0]
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def idea(number, title, oneliner, problem, solution_intro, code, features=None,
         tech=None, demo=None, strengths=None, risks=None, extra=None):
    doc.add_heading(f"{number}. {title}", level=2)
    para(oneliner, bold_prefix="One-liner: ", italic=True)
    para(problem, bold_prefix="Problem: ")
    para(solution_intro, bold_prefix="Solution: ")
    if code:
        codeblock(code)
    if features:
        para("", bold_prefix="Key features:", space_after=2)
        for f in features:
            bullet(f)
    if tech:
        para(tech, bold_prefix="Tech: ")
    if demo:
        para(demo, bold_prefix="Demo arc (2 min): ")
    if extra:
        para(extra)
    if strengths:
        para(strengths, bold_prefix="Strengths: ")
    if risks:
        para(risks, bold_prefix="Risks: ")


# ---------------- Title page ----------------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("BuildSprint Idea Briefs")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = ACCENT

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Prepared for team discussion — LatentForce.ai BuildSprint\n(48-hour online build sprint, no fixed theme)")
r.font.size = Pt(13)
r.font.color.rgb = GREY

para()
para("Nine candidate ideas, grouped into three families:", bold_prefix="")
bullet("PromptCI, AgentProfiler, DeadChunk, BreakBot", bold_prefix="Family A — AI-Infrastructure / Dev Tools: ")
bullet("ProofOfDelete, ConsentLedger, PPDA", bold_prefix="Family B — Privacy & Trust Products: ")
bullet("GitGuard, CodeTwin", bold_prefix="Family C — Code Provenance: ")

# ---------------- Family A ----------------
doc.add_page_break()
doc.add_heading("Family A — AI-Infrastructure / Dev Tools", level=1)
para("Theme: everyone is shipping LLM systems, but the tooling to test, debug, profile, and maintain them barely exists. These projects fill gaps in the AI field itself.", italic=True)

idea(
    1, "PromptCI — Regression testing for prompts  (AI-infra pick #1)",
    '"Prompts are code. Nobody tests them. We built the CI."',
    "Teams change prompts, bump model versions, or switch providers to save cost — and discover what broke from angry users. Prompt changes are the most-deployed, least-tested artifact in modern software.",
    "A CI pipeline for prompts. Save real inputs/outputs from logs as test cases. On every prompt/model change, PromptCI re-runs all cases against OLD vs NEW and reports:",
    ["\u2713 34 cases unchanged",
     "\u26A0 3 cases changed semantically (side-by-side diff)",
     "\u2717 2 cases now fail assertions:",
     "   - case #12: output no longer valid JSON",
     "   - case #29: total_amount wrong (\u20B94,820 \u2192 \u20B9482)",
     "\u2192 Block merge  /  Approve & re-baseline"],
    features=[
        "Deterministic assertions (JSON schema, regex, must-contain, numeric tolerance) + semantic diff for the rest",
        "\u201CCapture from logs \u2192 save as test case\u201D workflow (test creation is nearly free)",
        "Baseline-approval flow wired into git (block merge on regression)",
    ],
    tech="CLI + FastAPI backend, test runner, embedding-based semantic diff, simple web report UI.",
    demo="Edit one line of a prompt \u2192 push \u2192 report catches a regression the naked eye would miss \u2192 block merge \u2192 fix \u2192 green.",
    strengths="Maximum resonance with AI-dev-tools judges; clean 48h decomposition; real unowned gap (as a simple, self-hostable git-native tool).",
    risks="Space is warming up (LangSmith, Braintrust adjacent) — differentiation must live in the workflow, not the concept.",
)

idea(
    2, "AgentProfiler — The flame graph for LLM pipelines",
    '"Where do your agent\u2019s tokens, seconds, and rupees actually go?"',
    "An agent run costs \u20B918 and takes 40s. Why? Which step? Which retries? What % of the context window is boilerplate re-sent 12 times? Today the answer is \u201Cgrep the logs.\u201D Normal software has profilers/APM; AI has a void.",
    "A lightweight Python decorator/proxy that wraps LLM calls and renders a cost/latency/token flame graph per run:",
    ["Run #4412 \u2014 \u20B918.40, 41.2s, 9 LLM calls",
     "\u2588\u2588\u2588\u2588\u2588\u2588\u2588\u2588\u2588\u2588\u2588\u2588\u2588 planner   41% (\u20B97.60) \u2190 3 retries!",
     "\u26A0 61% of input tokens = repeated system prompt",
     "\u2192 Cache-hit potential: \u20B96.10/run (33%)"],
    features=[
        "Per-step cost/latency breakdown",
        "Retry detection",
        "Repeated-context waste analysis",
        "Cache-savings estimator",
    ],
    tech="100% deterministic engineering — token counting, timing, tree visualization. No AI needed in the build; nothing can flake mid-demo.",
    demo="Run an agent \u2192 open the flame graph \u2192 spot the waste \u2192 apply fix \u2192 re-run at 33% lower cost.",
    strengths="Zero-flake demo; unglamorous space = little competition; judge-proof punchline.",
    risks="Needs a realistic demo agent to profile; less \u201Cwow\u201D for non-technical audiences.",
)

idea(
    3, "DeadChunk — Auditor for rotting RAG knowledge bases",
    '"Your RAG answers are only as good as your knowledge base — and nobody audits the knowledge base."',
    "RAG knowledge bases rot silently: chunks that are never retrieved, near-duplicate chunks that split relevance scores, and chunks that contradict each other (old policy vs new policy). Teams debug the retriever and the prompt, never the data.",
    "Point it at a vector store + query logs. It reports:",
    ["\u2717 412 chunks never retrieved (dead weight, 31% of KB)",
     "\u2717 18 contradiction pairs (\u201Crefund within 30 days\u201D vs \u201C14 days\u201D)",
     "\u26A0 96 near-duplicate clusters splitting retrieval scores",
     "\u26A0 44 stale chunks (superseded by newer docs)",
     "\u2192 Cleaned KB export + before/after retrieval comparison"],
    tech="Embeddings + clustering for duplicates; NLI or LLM-judge for contradictions; log analysis for dead chunks.",
    demo="Ask a question \u2192 RAG gives a wrong answer sourced from a stale chunk \u2192 DeadChunk flags the contradiction \u2192 clean \u2192 same question now answered correctly.",
    strengths="Genuinely under-served gap; strong \u201Caha\u201D demo (contradictory chunks are shocking).",
    risks="Medium — needs a convincing demo KB with planted rot; contradiction detection can false-positive.",
)

idea(
    4, "BreakBot — Dependency-upgrade impact analysis",
    '"npm audit tells you what\u2019s vulnerable. Nothing tells you what breaks."',
    "Everyone freezes dependencies at old versions because upgrading is terrifying. Dependabot opens the PR, CI goes red, you close it. The info needed to upgrade safely exists (changelogs, migration guides, your call sites) — no human wants to cross-reference them.",
    "Give it your repo + \u201Cupgrade pydantic 1.10 \u2192 2.8\u201D. It fetches the changelog, finds every place YOUR code touches changed APIs, and produces a verdict per call site:",
    ["\u2717 app/models.py:41    .dict() removed \u2192 use .model_dump()",
     "\u2717 app/api/user.py:88  validator syntax changed",
     "\u26A0 app/config.py:12    default coercion changed silently",
     "\u2713 61 other usages unaffected",
     "[Generate fix PR] \u2192 patched diff, ready to review"],
    features=[
        "Every verdict cites the changelog lines it is based on (credibility feature)",
        "AI is essential here — changelog-to-callsite reasoning cannot be done with regex",
    ],
    tech="Repo ingestion (AST parsing for imports/usages), changelog fetcher, LLM verdict per call site, patch generation, web report.",
    demo="Red impact report \u2192 click \u201CGenerate fix PR\u201D \u2192 tests go green. Scope honestly: one language (Python), 3\u20134 famous breaking upgrades (pydantic 1\u21922, SQLAlchemy 1.4\u21922.0).",
    strengths="Dev judges feel this pain personally; agentic AI doing real reasoning, not a chatbot wrapper.",
    risks="Patch generation quality; must pre-select demo repos/upgrades carefully.",
)

# ---------------- Family B ----------------
doc.add_page_break()
doc.add_heading("Family B — Privacy & Trust Products", level=1)
para("Theme: leverages our existing PPDA codebase (hash-chained audit logs, AES-GCM, JWT, Shamir secret sharing, ONNX on-device inference). Roughly 30\u201340% head start on infrastructure.", italic=True)

idea(
    5, "ConsentLedger — \u201CA receipt for your data\u201D",
    '"Companies keep consent records for legal protection. Users keep nothing."',
    "When you sign up for an app, upload a document, or talk to an AI, you have no record of what you consented to, what data left your device, or who can see it.",
    "A personal, tamper-evident ledger of every data-sharing event:",
    ["Event: Uploaded resume to JobPortal.com",
     "Data: name, phone, employment history",
     "Consent scope: \u201Crecruitment matching only\u201D",
     "Retention claimed: 180 days",
     "Hash-chained record \u2713 (provably un-editable)",
     "Later: \u201CJobPortal\u2019s policy changed on 12 Mar \u2014",
     "        no longer matches your stored consent. Review?\u201D"],
    features=[
        "Hash-chained (tamper-evident) event log",
        "AES-GCM encrypted payloads",
        "Policy-text differ that alerts when terms drift from recorded consent",
    ],
    tech="Reuses PPDA\u2019s audit hash-chain + encryption-at-rest + JWT auth nearly as-is. Add capture flow (extension or upload) + dashboard + policy differ.",
    demo="Record consents \u2192 live-tamper with the database \u2192 ledger screams \u2192 policy-diff alert fires. Theatrical and 10 seconds to show.",
    strengths="~40% pre-built; India\u2019s DPDP Act makes consent records legally hot right now; \u201Creceipt for your data\u201D lands in 20 seconds.",
    risks="Consumer adoption story is hand-wavy (who installs this?); capture flow must not look like vaporware.",
)

idea(
    6, "ProofOfDelete — Verifiable data-deletion certificates",
    '"They said \u2018your data has been deleted.\u2019 Prove it."',
    "GDPR/DPDP grant the right to be forgotten, but when a company claims deletion, users just have to believe them. There is no verifiable artifact.",
    "A deletion-attestation service. Companies integrate a small SDK that emits a cryptographically signed deletion certificate (record hash + timestamp + position in a public hash chain). Users hold verifiable proof; regulators get an audit trail.",
    None,
    tech="Signing service + hash chain (reuse PPDA audit chain), SDK stub, verification portal.",
    demo="Simulate a company deleting a record \u2192 certificate issued \u2192 user verifies \u2192 then simulate a FAKE deletion \u2192 verification fails.",
    strengths="Sharp B2B compliance story (DPDP Act tailwind); technically distinctive.",
    risks="Two-sided market — demo requires simulating the company side; judges may ask \u201Cwhy would companies adopt this?\u201D (answer: regulatory pressure, but it\u2019s a debate).",
)

idea(
    7, "PPDA — Privacy-Preserving Digital Assistant (existing project)",
    '"A personal assistant where the privacy guarantees are cryptographic, not promises."',
    "Cloud assistants require trusting the provider with calendars, reminders, and private messages. PPDA is a local-first assistant backend where privacy is enforced by cryptography.",
    "Already built — FastAPI + PostgreSQL/SQLite backend with:",
    None,
    features=[
        "Zero-trust JWT auth, rate limiting, IDOR immunity",
        "AES-256-GCM encryption at rest for events/reminders/messages",
        "On-device ONNX intent classifier (<1 ms) + occlusion-saliency explainability",
        "Tamper-evident audit logging (append-only triggers + SHA-256 hash chaining)",
        "Cryptographic federated learning: Bonawitz secure aggregation (X25519, ChaCha20 masking, Shamir threshold shares) + R\u00E9nyi DP accounting",
    ],
    extra="As a BuildSprint entry: strongest technical depth of all nine, but it is pre-existing work (sprint rules and judge perception may penalize that), and \u201Cprivacy assistant\u201D takes longer than 20 seconds to pitch. Best used as an infrastructure quarry for ConsentLedger / ProofOfDelete rather than as the entry itself.",
)

# ---------------- Family C ----------------
doc.add_page_break()
doc.add_heading("Family C — Code Provenance", level=1)
para("Theme: \u201Cwas this project derived from that one?\u201D Discussed earlier; kept here for completeness with our earlier verdicts.", italic=True)

idea(
    8, "GitGuard — Explainable project provenance",
    '"Code similarity says \u2018these look alike.\u2019 GitGuard asks \u2018is there evidence this project was derived from that registered project?\u2019"',
    "Plain code-similarity tools cannot answer the provenance question: whether a later repository was derived from an earlier, registered project.",
    "Register a project \u2192 generate a digital fingerprint (code, AST, dependencies, structure, semantics, metadata/timestamps) \u2192 compare future repositories against it \u2192 output an explainable provenance score with an evidence breakdown:",
    ["Registered Project \u2192 Digital Fingerprint \u2192 Future Repository",
     "Evidence: code sim | AST sim | dependency sim |",
     "          structure | semantics | provenance metadata",
     "\u2192 Explainable Provenance Score",
     "\u2192 Original / Related / Potentially Derived / Highly Similar"],
    extra="Differentiator vs plain similarity tools: similarity is EVIDENCE, not the product. The product is the registration \u2192 fingerprint \u2192 later comparison \u2192 explainable evidence workflow.",
    strengths="Broader problem statement than clone detection; explainability angle.",
    risks="Judges may still bucket it as \u201Canother plagiarism detector\u201D; provenance claims are hard to validate rigorously in 48h; weakest fit for an AI-dev-tools audience among the strong candidates.",
)

idea(
    9, "CodeTwin — Structural/semantic code similarity  (previously rejected \u26A0)",
    '"Find structurally and semantically similar code across repositories."',
    "Detecting similar/cloned code across repositories.",
    "AST normalization + code fingerprints + semantic embeddings \u2192 similarity scores between codebases.",
    None,
    extra="Earlier verdict (unchanged): DO NOT build as a standalone project. It overlaps with existing code-similarity tools (e.g. the \u201CPonytail\u201D comparison) and reads as \u201Canother code clone detector.\u201D Its techniques are better used as the similarity-evidence engine inside GitGuard — or dropped entirely.",
)

# ---------------- Comparison matrix ----------------
doc.add_page_break()
doc.add_heading("Comparison Matrix", level=1)

headers = ["Idea", "Family", "Problem clarity (20s test)", "AI-judge appeal", "48h feasibility", "Novelty", "Uses our skills/code", "Overall"]
rows = [
    ["PromptCI", "AI-infra", "5/5", "5/5", "5/5", "4/5", "FastAPI, LLM orchestration", "#1"],
    ["AgentProfiler", "AI-infra", "4/5", "5/5", "5/5", "5/5", "Pure engineering", "#2"],
    ["BreakBot", "AI-infra", "5/5", "4/5", "4/5", "4/5", "Python/AST, LLM", "#3"],
    ["DeadChunk", "AI-infra", "4/5", "4/5", "3/5", "4/5", "Embeddings, RAG experience", "Strong"],
    ["ConsentLedger", "Privacy", "5/5", "3/5", "4/5", "4/5", "~40% from PPDA", "Strong"],
    ["ProofOfDelete", "Privacy", "4/5", "3/5", "3/5", "4/5", "Hash chain from PPDA", "Medium"],
    ["PPDA", "Privacy", "2/5", "3/5", "(pre-built)", "3/5", "It IS our code", "Quarry, not entry"],
    ["GitGuard", "Provenance", "3/5", "2/5", "3/5", "3/5", "Some", "Backup only"],
    ["CodeTwin", "Provenance", "3/5", "1/5", "3/5", "1/5", "Some", "Rejected"],
]

table = doc.add_table(rows=1 + len(rows), cols=len(headers))
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        cell = table.rows[i + 1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                if j == 0:
                    r.bold = True

para()
doc.add_heading("Recommended team decision framing", level=1)
bullet("PromptCI (max resonance) or AgentProfiler (zero-flake demo, least competition).",
       bold_prefix="If judges are AI-dev-tools people (LatentForce = yes): pick from Family A \u2192 ")
bullet("ConsentLedger (PPDA hash-chain + crypto, DPDP Act tailwind).",
       bold_prefix="If we want to reuse the most existing code: ")
bullet("BreakBot.", bold_prefix="If we want AI doing visible heavy-lifting in the demo: ")
bullet("CodeTwin standalone; PPDA as-is (pre-existing work); GitGuard only as fallback.",
       bold_prefix="Avoid: ")
para()
p = para("Suggested vote: each teammate ranks their top 3; tie-break on \u201Cwhich demo can we make bulletproof by hour 40?\u201D")
for r in p.runs:
    r.bold = True

doc.save("/home/user/Privacy-Preserving-PDA/docs/BuildSprint-Idea-Briefs.docx")
print("saved")
